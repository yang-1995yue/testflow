"""
文件内容提取工具
支持从不同格式的文件中提取文本内容和图片
"""
import os
from typing import Optional, List, Dict, Any
from pathlib import Path


def extract_text_from_file(file_path: str, file_type: str) -> tuple[str, Optional[str]]:
    """
    从文件中提取文本内容
    
    Args:
        file_path: 文件路径
        file_type: 文件类型 (txt/docx/md)
        
    Returns:
        tuple[content, error]: (提取的内容, 错误信息)
    """
    try:
        if file_type == 'txt':
            return extract_from_txt(file_path)
        elif file_type == 'docx':
            return extract_from_docx(file_path)
        elif file_type == 'md':
            return extract_from_md(file_path)
        else:
            return "", f"不支持的文件类型: {file_type}"
    except Exception as e:
        return "", f"提取失败: {str(e)}"


def extract_from_txt(file_path: str) -> tuple[str, Optional[str]]:
    """
    从TXT文件提取内容
    """
    try:
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                return content.strip(), None
            except UnicodeDecodeError:
                continue
        
        return "", "无法识别文件编码，请确保文件是有效的文本文件"
    except Exception as e:
        return "", f"读取TXT文件失败: {str(e)}"


def extract_from_docx(file_path: str) -> tuple[str, Optional[str]]:
    """
    从DOCX文件提取内容
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # 提取所有段落
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # 忽略空段落
                paragraphs.append(text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    paragraphs.append(' | '.join(row_data))
        
        content = '\n\n'.join(paragraphs)
        
        if not content:
            return "", "文档为空或无法提取内容"
        
        return content, None
    except ImportError:
        return "", "缺少python-docx库，请安装: pip install python-docx"
    except Exception as e:
        return "", f"读取DOCX文件失败: {str(e)}"


def extract_from_md(file_path: str) -> tuple[str, Optional[str]]:
    """
    从Markdown文件提取内容
    """
    try:
        # Markdown文件本质是文本文件，直接读取即可
        encodings = ['utf-8', 'gbk', 'gb2312']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                return content.strip(), None
            except UnicodeDecodeError:
                continue
        
        return "", "无法识别Markdown文件编码"
    except Exception as e:
        return "", f"读取Markdown文件失败: {str(e)}"


def validate_file_type(filename: str) -> tuple[bool, str]:
    """
    验证文件类型
    
    Args:
        filename: 文件名
        
    Returns:
        tuple[is_valid, file_type]: (是否有效, 文件类型)
    """
    ext = os.path.splitext(filename)[1].lower()
    
    allowed_extensions = {
        '.txt': 'txt',
        '.docx': 'docx',
        '.md': 'md'
    }
    
    if ext in allowed_extensions:
        return True, allowed_extensions[ext]
    else:
        return False, ""


def validate_file_size(file_size: int, max_size: int = 10 * 1024 * 1024) -> bool:
    """
    验证文件大小
    
    Args:
        file_size: 文件大小（字节）
        max_size: 最大允许大小（字节），默认10MB
        
    Returns:
        bool: 是否有效
    """
    return file_size <= max_size


def extract_images_from_docx(file_path: str, output_dir: str) -> tuple[List[Dict[str, Any]], Optional[str]]:
    """
    从DOCX文件中提取所有嵌入图片
    
    Args:
        file_path: DOCX文件路径
        output_dir: 图片输出目录路径
        
    Returns:
        tuple[images, error]: (图片信息列表, 错误信息)
        
        图片信息字典包含:
        - path: 图片保存路径
        - format: 图片格式 (png/jpg/gif等)
        - size: 图片大小（字节）
        - position_index: 在文档中的位置顺序
        - width: 图片宽度（如果可获取）
        - height: 图片高度（如果可获取）
    """
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        return [], "缺少python-docx库，请安装: pip install python-docx"
    
    try:
        # 确保输出目录存在
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        doc = Document(file_path)
        images: List[Dict[str, Any]] = []
        
        # 遍历文档中的所有关系，找到图片
        for rel_id, rel in doc.part.rels.items():
            # 检查是否是图片关系
            if "image" in rel.target_ref.lower():
                try:
                    image_data = rel.target_part.blob
                    
                    # 从target_ref中提取文件扩展名
                    target_ref = rel.target_ref
                    image_ext = target_ref.split('.')[-1].lower() if '.' in target_ref else 'png'
                    
                    # 标准化扩展名
                    if image_ext == 'jpeg':
                        image_ext = 'jpg'
                    elif image_ext not in ['png', 'jpg', 'gif', 'bmp', 'tiff', 'webp']:
                        image_ext = 'png'  # 默认使用png
                    
                    # 生成唯一文件名
                    position_index = len(images)
                    image_filename = f"image_{position_index:03d}.{image_ext}"
                    image_path = output_path / image_filename
                    
                    # 保存图片
                    with open(image_path, 'wb') as f:
                        f.write(image_data)
                    
                    # 尝试获取图片尺寸
                    width, height = _get_image_dimensions(str(image_path))
                    
                    images.append({
                        'path': str(image_path),
                        'format': image_ext,
                        'size': len(image_data),
                        'position_index': position_index,
                        'width': width,
                        'height': height
                    })
                except Exception as e:
                    # 单个图片提取失败不影响其他图片
                    print(f"提取图片失败 (rel_id={rel_id}): {str(e)}")
                    continue
        
        return images, None
        
    except Exception as e:
        return [], f"提取DOCX图片失败: {str(e)}"


def _get_image_dimensions(image_path: str) -> tuple[Optional[int], Optional[int]]:
    """
    获取图片尺寸
    
    Args:
        image_path: 图片文件路径
        
    Returns:
        tuple[width, height]: (宽度, 高度)，如果无法获取则返回(None, None)
    """
    try:
        # 尝试使用PIL获取图片尺寸
        from PIL import Image
        with Image.open(image_path) as img:
            return img.size
    except ImportError:
        # PIL未安装，尝试读取PNG/JPEG头部信息
        return _get_dimensions_from_header(image_path)
    except Exception:
        return None, None


def _get_dimensions_from_header(image_path: str) -> tuple[Optional[int], Optional[int]]:
    """
    从图片文件头部读取尺寸信息（不依赖PIL）
    
    支持PNG和JPEG格式
    """
    try:
        with open(image_path, 'rb') as f:
            header = f.read(32)
            
            # PNG格式
            if header[:8] == b'\x89PNG\r\n\x1a\n':
                # PNG的IHDR块在第16-24字节包含宽高
                width = int.from_bytes(header[16:20], 'big')
                height = int.from_bytes(header[20:24], 'big')
                return width, height
            
            # JPEG格式
            if header[:2] == b'\xff\xd8':
                f.seek(0)
                f.read(2)  # 跳过SOI标记
                
                while True:
                    marker = f.read(2)
                    if len(marker) < 2:
                        break
                    
                    if marker[0] != 0xff:
                        break
                    
                    # SOF0, SOF1, SOF2 标记包含尺寸信息
                    if marker[1] in [0xc0, 0xc1, 0xc2]:
                        f.read(3)  # 跳过长度和精度
                        height = int.from_bytes(f.read(2), 'big')
                        width = int.from_bytes(f.read(2), 'big')
                        return width, height
                    
                    # 跳过其他段
                    if marker[1] != 0xd8 and marker[1] != 0xd9:
                        length = int.from_bytes(f.read(2), 'big')
                        f.seek(length - 2, 1)
        
        return None, None
    except Exception:
        return None, None
